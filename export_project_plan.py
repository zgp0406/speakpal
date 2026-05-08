from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, align=None, first_indent=0, space_before=0, space_after=0, line_spacing=1.4):
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(first_indent)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_run(paragraph, text, font_name="宋体", size=10.5, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return run


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

styles = doc.styles
styles["Normal"].font.name = "宋体"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(10.5)

# Cover
p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=120, space_after=18, line_spacing=1.0)
add_run(p, "项目实施与系统设计规划", font_name="SimHei", size=24, bold=True, color="1F4E79")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=16, line_spacing=1.1)
add_run(p, "面向儿童英语发音的数字人口型与舌位可视化纠错系统", font_name="SimHei", size=16, bold=True)

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6, line_spacing=1.0)
add_run(p, "适用于开题后启动开发阶段的实施规划文档", font_name="宋体", size=11, color="666666")

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
widths = [Cm(4.0), Cm(10.5)]
for left, right in [
    ("文档内容", "项目实施计划、功能清单、系统架构、数据库设计、任务拆分"),
    ("适用阶段", "正式开发前规划与分工"),
    ("开发平台", "Android + 后端服务 + 数字人展示"),
    ("输出形式", "Word 文档"),
]:
    row = meta.add_row().cells
    row[0].width = widths[0]
    row[1].width = widths[1]
    for c in row:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, 120, 140, 120, 140)
    shade_cell(row[0], "D9EAF7")
    p1 = row[0].paragraphs[0]
    format_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p1, left, font_name="SimHei", size=10.5, bold=True)
    p2 = row[1].paragraphs[0]
    format_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    add_run(p2, right, font_name="宋体", size=10.5)

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=8, line_spacing=1.0)
add_run(p, "2026年5月", font_name="宋体", size=11)

doc.add_page_break()

def add_heading(text):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=4, space_after=4, line_spacing=1.0)
    add_run(p, text, font_name="SimHei", size=12, bold=True, color="1F4E79")

def add_body(text, indent=0.74):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=indent, space_before=0, space_after=3, line_spacing=1.35)
    add_run(p, text, font_name="宋体", size=10.5)

add_heading("一、项目实施计划表")
plan = doc.add_table(rows=1, cols=4)
plan.style = "Table Grid"
plan.alignment = WD_TABLE_ALIGNMENT.CENTER
plan.autofit = False
plan_widths = [Cm(2.4), Cm(2.4), Cm(4.6), Cm(5.1)]
headers = ["阶段", "时间", "目标", "主要产出"]
hdr = plan.rows[0].cells
for i, w in enumerate(plan_widths):
    hdr[i].width = w
    hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(hdr[i], 120, 120, 120, 120)
    shade_cell(hdr[i], "D9EAF7")
    p = hdr[i].paragraphs[0]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p, headers[i], font_name="SimHei", size=10.5, bold=True)

plan_rows = [
    ("需求分析", "第1周", "明确用户、功能范围、技术路线", "需求说明、功能清单、原型草图"),
    ("总体设计", "第2周", "完成系统架构、数据库、页面结构设计", "架构图、数据库表、页面流程图"),
    ("Android基础开发", "第3周", "完成首页、练习页、结果页骨架", "页面导航、基础UI、主题样式"),
    ("录音与播放", "第4周", "实现标准发音播放和用户录音", "音频采集、播放控制、录音保存"),
    ("评测接入", "第5周", "接入语音评测接口并解析结果", "评分、错误音素、反馈信息"),
    ("纠错展示", "第6周", "生成适合儿童的纠错提示", "错误提示文案、结果页展示"),
    ("数字人可视化", "第7周", "实现口型与舌位演示", "口型动画、舌位示意、同步播放"),
    ("学习记录", "第8周", "保存练习数据并展示统计", "历史记录、统计页、错音汇总"),
    ("联调测试", "第9周", "修复 bug，优化交互体验", "测试报告、问题修复记录"),
    ("论文与答辩", "第10周", "整理材料、截图、文档", "论文初稿、PPT、演示视频"),
]
for row_data in plan_rows:
    cells = plan.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].width = plan_widths[i]
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[i], 120, 120, 120, 120)
        p = cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.1)
        add_run(p, val, font_name="宋体", size=10 if i >= 2 else 10.5, bold=(i == 0))

add_heading("二、详细功能清单")
feature_sections = {
    "用户端功能": ["首页", "单词列表页", "单词详情页", "跟读练习页", "评测结果页", "数字人演示页", "学习记录页", "我的统计页"],
    "核心训练功能": ["播放标准发音", "录制用户发音", "重新录音", "语音上传评测", "获取评分结果", "展示错误音素", "展示纠错建议", "重复练习同一单词"],
    "数字人可视化功能": ["标准口型展示", "舌位展示", "慢速播放", "分步演示", "当前音素高亮", "动作提示文字", "鼓励性反馈动画"],
    "学习管理功能": ["练习次数统计", "历史成绩记录", "常错音素统计", "单词难度分类", "今日练习完成度", "学习进度曲线"],
    "系统管理功能": ["单词数据管理", "音标与规则配置", "标准音频资源管理", "评测结果缓存", "用户训练记录存储"],
}
for title, items in feature_sections.items():
    p = doc.add_paragraph()
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    add_run(p, title, font_name="SimHei", size=10.8, bold=True)
    add_body("；".join(items) + "。", indent=0.74)

add_heading("三、系统架构图")
for line in [
    "Android客户端 -> 录音模块、音频播放模块、评测请求模块、数字人展示模块、本地存储模块",
    "后端服务 -> 语音评测、错误音素识别、纠错建议生成、成绩返回",
    "展示模块 -> 口型动画、舌位示意、发音步骤提示",
    "数据层 -> 单词库、学习记录、练习统计",
]:
    add_body(line, indent=0.74)
p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=4, line_spacing=1.0)
add_run(p, "可直接在论文中绘制为流程图，也可后续替换为 Mermaid/Visio 图。", font_name="宋体", size=9.5, color="666666")

add_heading("四、数据库表设计")
db_tables = [
    ("user", "id, username, age_group, created_at", "用户基本信息"),
    ("word", "id, word_text, phonetic, difficulty, category, audio_path, created_at", "单词与标准音频"),
    ("phoneme_rule", "id, phoneme, mouth_hint, tongue_hint, note", "发音规则说明"),
    ("practice_record", "id, user_id, word_id, score, result_level, error_phonemes, record_time", "练习记录"),
    ("evaluation_result", "id, record_id, raw_result, feedback_text, created_at", "评测结果"),
    ("learning_stats", "id, user_id, total_practice, avg_score, common_error, updated_at", "学习统计"),
]
db = doc.add_table(rows=1, cols=3)
db.style = "Table Grid"
db.alignment = WD_TABLE_ALIGNMENT.CENTER
db.autofit = False
db_widths = [Cm(3.2), Cm(8.0), Cm(4.0)]
for i, h in enumerate(["表名", "字段", "说明"]):
    db.rows[0].cells[i].width = db_widths[i]
    db.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(db.rows[0].cells[i], 120, 120, 120, 120)
    shade_cell(db.rows[0].cells[i], "D9EAF7")
    p = db.rows[0].cells[i].paragraphs[0]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p, h, font_name="SimHei", size=10.5, bold=True)
for row_data in db_tables:
    cells = db.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].width = db_widths[i]
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[i], 120, 120, 120, 120)
        p = cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if i != 0 else WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.1)
        add_run(p, val, font_name="宋体", size=10.0)

add_heading("五、前后端开发任务拆分")
tasks = {
    "Android前端任务": ["搭建项目基础结构", "设计首页和练习页", "实现单词列表", "实现标准发音播放", "实现录音按钮和录音保存", "实现评测结果页", "实现数字人展示页", "实现学习记录页", "接入本地数据库", "做页面跳转和状态管理"],
    "后端任务": ["提供音频上传接口", "提供语音评测接口", "返回评分和错误音素", "返回纠错建议", "保存练习记录", "查询历史成绩和统计数据"],
    "算法/评测任务": ["设计评分规则", "定义错误音素识别逻辑", "设计儿童友好的反馈模板", "处理单词级发音结果", "将原始结果转成可读提示"],
    "数字人/可视化任务": ["设计口型动画资源", "设计舌位示意图", "做慢速播放和分步骤演示", "设计高亮提示", "设计鼓励反馈动画"],
    "测试任务": ["测试录音是否成功", "测试标准音频播放", "测试评测接口稳定性", "测试结果页展示是否正常", "测试数据库是否正确保存", "测试不同机型兼容性"],
}
for title, items in tasks.items():
    p = doc.add_paragraph()
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    add_run(p, title, font_name="SimHei", size=10.8, bold=True)
    add_body("；".join(items) + "。", indent=0.74)

doc.core_properties.author = "OpenAI"
doc.core_properties.title = "面向儿童英语发音的数字人口型与舌位可视化纠错系统 - 项目实施规划"
doc.core_properties.subject = "项目规划"
doc.core_properties.category = "Planning"

doc.save("项目规划导出.docx")
