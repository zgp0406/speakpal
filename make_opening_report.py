from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, align=None, first_indent=0, space_before=0, space_after=0, line_spacing=1.5):
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(first_indent)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_run(paragraph, text, font_name="宋体", size=10.5, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, name=font_name, size=size, bold=bold, color=color)
    return run


def add_field_run(paragraph, field_code, result_text=""):
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr)
    r.append(sep)
    if result_text:
        t = OxmlElement("w:t")
        t.text = result_text
        r.append(t)
    r.append(end)
    return run


def add_page_number_paragraph(container, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = container.add_paragraph()
    format_paragraph(p, align=align, space_before=0, space_after=0, line_spacing=1.0)
    add_run(p, "第 ", font_name="宋体", size=9, color="666666")
    add_field_run(p, "PAGE", "1")
    add_run(p, " 页", font_name="宋体", size=9, color="666666")
    return p


def add_header_footer(section, header_text="开题报告"):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header_p = header.paragraphs[0]
    header_p.clear()
    format_paragraph(header_p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, line_spacing=1.0)
    add_run(header_p, header_text, font_name="宋体", size=9, color="666666")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    format_paragraph(footer_p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, line_spacing=1.0)
    add_run(footer_p, "第 ", font_name="宋体", size=9, color="666666")
    add_field_run(footer_p, "PAGE", "1")
    add_run(footer_p, " 页", font_name="宋体", size=9, color="666666")


def set_section_layout(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


doc = Document()
doc.settings._element.append(OxmlElement("w:updateFields"))
doc.settings._element[-1].set(qn("w:val"), "true")

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(10.5)

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    s = styles[style_name]
    s.font.name = "SimHei"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

styles["Heading 1"].font.size = Pt(13)
styles["Heading 2"].font.size = Pt(11)
styles["Heading 3"].font.size = Pt(10.5)

cover = doc.sections[0]
set_section_layout(cover)
cover.different_first_page_header_footer = True
cover.header.is_linked_to_previous = False
cover.footer.is_linked_to_previous = False
cover.header.paragraphs[0].clear()
cover.footer.paragraphs[0].clear()

# Cover page
p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=150, space_after=16, line_spacing=1.0)
add_run(p, "开题报告", font_name="SimHei", size=24, bold=True, color="1F4E79")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=100, space_after=20, line_spacing=1.05)
add_run(p, "面向儿童英语发音的数字人口型与舌位可视化纠错系统", font_name="SimHei", size=18, bold=True)

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12, line_spacing=1.0)
add_run(p, "Android 平台英语发音辅助学习系统设计", font_name="宋体", size=11, color="666666")

info_rows = [
    ("课题名称", "面向儿童英语发音的数字人口型与舌位可视化纠错系统"),
    ("课题性质", "应用研究 / 系统设计"),
    ("开发平台", "Android"),
    ("完成形式", "软件系统 + 开题报告"),
]
tbl = doc.add_table(rows=0, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.autofit = False
widths = [Cm(4.0), Cm(10.5)]
for left, right in info_rows:
    row = tbl.add_row().cells
    row[0].width = widths[0]
    row[1].width = widths[1]
    for c in row:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, 120, 140, 120, 140)
    set_cell_shading(row[0], "D9EAF7")
    p1 = row[0].paragraphs[0]
    format_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p1, left, font_name="SimHei", size=10.5, bold=True)
    p2 = row[1].paragraphs[0]
    format_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    add_run(p2, right, font_name="宋体", size=10.5)

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=10, line_spacing=1.0)
add_run(p, "2026年5月", font_name="宋体", size=12)

# TOC section
toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
set_section_layout(toc_section)
add_header_footer(toc_section, "开题报告")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=10, line_spacing=1.0)
add_run(p, "目录", font_name="SimHei", size=18, bold=True, color="1F4E79")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2, line_spacing=1.0)
add_field_run(p, 'TOC \\o "1-1" \\h \\z \\u', "目录会在 Word 中自动更新")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=0, line_spacing=1.0)
add_run(p, "提示：如目录未刷新，在 Word 中全选后按 F9 更新字段。", font_name="宋体", size=9, color="666666")

# Body section
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
set_section_layout(body_section)
add_header_footer(body_section, "开题报告")

sections = [
    ("一、课题背景与研究意义", [
        "随着英语学习在基础教育阶段的重要性不断提高，儿童英语发音训练逐渐成为家长和教师关注的重点。然而，在实际学习过程中，儿童往往存在发音器官认知不足、发音位置不清楚、模仿能力有限等问题。传统英语发音教学主要依赖教师口头讲解、示范朗读和反复跟读，这种方式虽然能够在一定程度上帮助学生掌握发音，但缺乏直观的口型和舌位展示，儿童难以准确理解某些音素的发音机制，导致纠错效率不高，学习过程也相对枯燥。",
        "近年来，随着语音识别、语音评测和数字人技术的发展，将发音评测与可视化教学相结合，为英语发音学习提供了新的思路。语音评测技术可以对用户的发音进行分析与评分，识别发音偏差；数字人技术则可以通过动态展示口型变化和舌位运动，将抽象的发音过程直观呈现出来，使学习者更容易理解标准发音的动作要领。",
        "因此，研究并开发一套面向儿童英语发音学习的数字人口型与舌位可视化纠错系统，不仅具有较强的应用价值，也具有一定的研究意义。该系统能够将发音评测与可视化教学有机结合，为儿童英语发音训练提供更加直观、智能和高效的辅助工具。"
    ]),
    ("二、国内外研究现状", [
        "国外较早开始探索语音识别与发音训练结合的应用，并形成了一批基于移动端或在线平台的发音学习产品。这类系统通常具备标准发音播放、语音识别、发音评分、错误提示等功能，能够帮助学习者进行跟读训练。部分研究还尝试通过动画、口型图示或三维模型展示发音器官动作，以提升学习者对发音机制的理解。",
        "国内在英语发音辅助教学方面的研究起步较晚，但近年来发展较快。随着人工智能和语音技术的普及，越来越多的英语学习软件开始引入智能评测功能，如单词发音评分、语音对比分析和发音错误提示等。然而，从整体来看，现有系统大多偏重“评分”和“识别”，对“如何纠正发音”这一关键问题的支持仍不充分，特别是在儿童英语学习场景下，现有系统普遍存在反馈方式抽象、可视化表达单一、趣味性不足等问题。"
    ]),
    ("三、课题研究目标", [
        "本课题旨在设计并实现一套基于 Android 平台的儿童英语发音可视化纠错系统，实现发音采集、发音评测、口型与舌位可视化、纠错训练与学习记录等功能，并以直观、易懂的方式帮助儿童提升英语发音准确率。"
    ]),
    ("四、课题研究内容", [
        "1. 用户端语音采集与发音播放：实现儿童跟读录音功能，并提供标准单词发音播放，作为练习与对比的基础。",
        "2. 发音评测与错误识别：对儿童朗读的英文单词进行分析，得到发音评分及错误提示信息，判断整体发音准确度及可能存在的问题音素。",
        "3. 数字人口型与舌位可视化：在标准发音播放过程中，通过数字人动态展示口型变化、舌位位置和发音动作，将音标发音过程可视化。",
        "4. 可视化纠错反馈：根据评测结果，对错误发音进行针对性提示，辅助儿童调整发音。",
        "5. 练习记录与学习统计：记录每次练习的单词、评分、错误类型和练习时间，生成学习记录，便于用户查看学习进展。"
    ]),
    ("五、拟解决的关键问题", [
        "1. 如何将发音评测结果转化为适合儿童理解的纠错提示。",
        "2. 如何用数字人技术直观展示口型与舌位变化。",
        "3. 如何设计简单、友好的儿童交互界面。",
        "4. 如何构建评测、反馈、训练联动的学习闭环。",
        "5. 如何保证系统在 Android 平台上的稳定运行与较好的用户体验。"
    ]),
    ("六、研究方法与技术路线", [
        "本课题采用“需求分析 + 系统设计 + 原型实现 + 功能验证”的方法开展研究。首先分析儿童英语发音学习需求，明确系统功能；然后进行系统架构设计与数据库设计；接着实现 Android 端主要功能模块，并集成语音评测与数字人可视化展示；最后通过实际使用和测试验证系统效果。",
        "系统整体采用“Android 端 + 语音评测服务 + 数字人展示模块”的架构。用户在移动端完成发音录制后，音频被发送到评测模块进行分析，返回评分与错误信息；随后由可视化模块播放标准发音动作演示，并给出针对性纠错提示，形成完整的训练闭环。",
        "主要技术选型：Android 开发建议使用 Kotlin 或 Java；本地存储可采用 Room 或 SQLite；音频采集可使用 AudioRecord 或 MediaRecorder；音频播放可使用 ExoPlayer；后端服务可采用 Python FastAPI 或 Java Spring Boot；可视化展示可先采用二维动画，再逐步升级为三维数字人渲染。"
    ]),
    ("七、系统设计方案", [
        "系统可分为四层：表现层包括首页、单词练习页、评测结果页、数字人演示页和学习记录页；业务逻辑层负责录音管理、评测请求处理、错误提示生成与训练流程控制；服务能力层提供语音识别与评测、数字人动画播放、音频播放与录制、数据存储等能力；数据层则保存单词数据、音标与发音规则、学习记录和评测结果。",
        "核心功能流程为：用户选择单词后，系统播放标准发音并同步展示口型与舌位；用户录音跟读后，系统进行发音评测并返回分数和错误提示；系统根据评测结果生成纠错建议，用户可再次练习并查看学习记录。"
    ]),
    ("八、系统创新点", [
        "1. 面向儿童场景设计，注重界面友好性、交互趣味性和提示直观性。",
        "2. 发音评测与可视化纠错结合，不仅对发音进行评分，还通过数字人展示标准口型和舌位。",
        "3. 构建“听标准音-看发音动作-跟读练习-系统反馈-再次训练”的完整学习流程。",
        "4. 强调发音动作可视化，使抽象的发音机制更易于理解。"
    ]),
    ("九、可行性分析", [
        "技术可行性：Android 开发、音频采集、语音评测接口调用、数据存储与动画展示等技术均较成熟，具备实现基础。",
        "需求可行性：儿童英语发音学习存在明确需求，家长和教师对发音纠错、直观示范和趣味训练有较强接受度。",
        "实现可行性：本课题功能模块划分清晰，开发难度可控。可先实现单词级评测与二维口型演示，再逐步扩展数字人口型和舌位展示。"
    ]),
    ("十、预期成果", [
        "1. 完成一套基于 Android 平台的儿童英语发音可视化纠错系统。",
        "2. 实现语音采集、标准发音播放、发音评测、可视化纠错和学习记录等功能。",
        "3. 形成系统设计文档、数据库设计文档和测试报告。",
        "4. 完成答辩演示材料及相关论文内容。"
    ]),
    ("十一、进度安排", []),
    ("十二、参考文献", []),
]

for title, paras in sections:
    h = doc.add_paragraph()
    h.style = styles["Heading 1"]
    format_paragraph(h, space_before=6, space_after=4, line_spacing=1.0)
    add_run(h, title, font_name="SimHei", size=12, bold=True, color="1F4E79")
    if title == "十一、进度安排":
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        col_widths = [Cm(4.0), Cm(10.5)]
        hdr[0].width = col_widths[0]
        hdr[1].width = col_widths[1]
        for c in hdr:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c, 120, 140, 120, 140)
            set_cell_shading(c, "D9EAF7")
        set_repeat_table_header(table.rows[0])
        p0 = hdr[0].paragraphs[0]
        format_paragraph(p0, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
        add_run(p0, "阶段", font_name="SimHei", size=10.5, bold=True)
        p1 = hdr[1].paragraphs[0]
        format_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
        add_run(p1, "主要任务", font_name="SimHei", size=10.5, bold=True)

        plan_rows = [
            ("第一阶段", "需求分析与资料调研，明确系统目标、功能需求和实现方案。"),
            ("第二阶段", "系统设计，包括架构设计、数据库设计、界面原型设计和技术路线确定。"),
            ("第三阶段", "核心功能开发，实现录音、播放、评测调用、结果展示和记录存储等基础功能。"),
            ("第四阶段", "数字人可视化模块开发，完善口型和舌位动态展示效果。"),
            ("第五阶段", "系统测试与优化，对功能、稳定性和界面体验进行修正。"),
            ("第六阶段", "论文整理与答辩准备，完成论文撰写和答辩材料制作。"),
        ]
        for a, b in plan_rows:
            row = table.add_row().cells
            row[0].width = col_widths[0]
            row[1].width = col_widths[1]
            for c in row:
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(c, 120, 140, 120, 140)
            p = row[0].paragraphs[0]
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            add_run(p, a, font_name="宋体", size=10.5, bold=True)
            p = row[1].paragraphs[0]
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.2)
            add_run(p, b, font_name="宋体", size=10.0)
        continue

    if title == "十二、参考文献":
        refs = [
            "1. 语音识别与语音评测相关研究文献。",
            "2. 儿童英语发音教学与纠错研究文献。",
            "3. Android 移动应用开发相关文献。",
            "4. 数字人技术与可视化交互相关文献。",
            "5. 语音学习系统设计与实现相关论文。",
        ]
        for r in refs:
            p = doc.add_paragraph()
            format_paragraph(p, first_indent=0.74, space_before=0, space_after=2, line_spacing=1.3)
            add_run(p, r, font_name="宋体", size=10.5)
        continue

    for para in paras:
        p = doc.add_paragraph()
        format_paragraph(p, first_indent=0.74, space_before=0, space_after=3, line_spacing=1.4)
        add_run(p, para, font_name="宋体", size=10.5)

doc.core_properties.author = "OpenAI"
doc.core_properties.title = "面向儿童英语发音的数字人口型与舌位可视化纠错系统开题报告"
doc.core_properties.subject = "开题报告"
doc.core_properties.category = "Research Proposal"

doc.save("opening_report.docx")
