from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, align=None, first_indent=0, space_before=0, space_after=0, line_spacing=1.35):
    if align is not None:
        paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(first_indent)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_run(paragraph, text, font_name="宋体", size=10.5, bold=False, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return run


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.8)

styles = doc.styles
styles["Normal"].font.name = "宋体"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(10.5)


def title(text, size=16):
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8, line_spacing=1.0)
    add_run(p, text, font_name="SimHei", size=size, bold=True, color="1F4E79")


def body(text, indent=0.74, size=10.5):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=indent, space_before=0, space_after=3, line_spacing=1.35)
    add_run(p, text, font_name="宋体", size=size)


def bullet(text):
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=0.0, space_before=0, space_after=2, line_spacing=1.25)
    add_run(p, "• " + text, font_name="宋体", size=10.5)


# Cover
p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=120, space_after=16, line_spacing=1.0)
add_run(p, "项目落地蓝图", font_name="SimHei", size=24, bold=True, color="1F4E79")

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=18, line_spacing=1.1)
add_run(p, "面向儿童英语发音的数字人口型与舌位可视化纠错系统", font_name="SimHei", size=16, bold=True)

p = doc.add_paragraph()
format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10, line_spacing=1.0)
add_run(p, "基于开题报告与项目规划文件整理", font_name="宋体", size=11, color="666666")

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
for left, right in [
    ("定位", "开发前的技术实施蓝图"),
    ("目标", "让项目可以直接进入编码阶段"),
    ("范围", "Android前端、后端评测、数字人展示、数据管理"),
    ("版本", "MVP优先"),
]:
    row = meta.add_row().cells
    for i, c in enumerate(row):
        c.width = Cm(4.2 if i == 0 else 10.2)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, 120, 140, 120, 140)
    shade_cell(row[0], "D9EAF7")
    p1 = row[0].paragraphs[0]
    format_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p1, left, font_name="SimHei", size=10.5, bold=True)
    p2 = row[1].paragraphs[0]
    format_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
    add_run(p2, right, font_name="宋体", size=10.5)

doc.add_page_break()

title("一、项目定位与MVP边界", 13)
body("本项目面向儿童英语发音学习，重点解决“听不出差异、看不懂动作、改不准发音”三个问题。系统不追求一次性完成复杂的实时3D驱动，而是优先建立可用的发音训练闭环。")
body("MVP阶段建议只做单词级训练，流程固定为：选择单词 -> 播放标准音 -> 展示口型/舌位 -> 录音跟读 -> 评测打分 -> 返回纠错建议 -> 再次练习。")
bullet("优先级最高：录音、标准音播放、评分返回、错音提示。")
bullet("第二优先级：数字人口型与舌位可视化。")
bullet("第三优先级：学习记录、统计和成长曲线。")

title("二、功能模块拆分", 13)
module_sections = {
    "1. 学习入口模块": ["首页", "年级/难度入口", "推荐单词列表", "今日练习入口"],
    "2. 发音训练模块": ["标准音频播放", "录音采集", "暂停/重录", "播放-跟读-对比流程"],
    "3. 发音评测模块": ["上传音频", "获取分数", "识别错误音素", "生成儿童可读反馈"],
    "4. 可视化纠错模块": ["口型动画", "舌位示意", "音素高亮", "慢速分步展示"],
    "5. 学习管理模块": ["练习记录", "分数统计", "错音统计", "历史回放"],
}
for sec_title, items in module_sections.items():
    p = doc.add_paragraph()
    format_paragraph(p, space_before=2, space_after=2, line_spacing=1.0)
    add_run(p, sec_title, font_name="SimHei", size=11, bold=True)
    body("；".join(items) + "。", indent=0.74)

title("三、页面结构与交互流程", 13)
page_rows = [
    ("首页", "进入练习、查看进度、继续上次任务"),
    ("单词列表页", "按难度/音素分类选择练习内容"),
    ("练习页", "播放标准音、展示数字人、录音跟读"),
    ("结果页", "展示分数、错音、纠错建议、重新练习"),
    ("记录页", "查看历史成绩、常错音素、学习趋势"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
headers = ["页面", "核心动作"]
widths = [Cm(4.0), Cm(10.4)]
for i, h in enumerate(headers):
    c = table.rows[0].cells[i]
    c.width = widths[i]
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(c, 120, 120, 120, 120)
    shade_cell(c, "D9EAF7")
    p = c.paragraphs[0]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p, h, font_name="SimHei", size=10.5, bold=True)
for a, b in page_rows:
    cells = table.add_row().cells
    for i, text in enumerate([a, b]):
        cells[i].width = widths[i]
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[i], 120, 120, 120, 120)
        p = cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.05)
        add_run(p, text, font_name="宋体", size=10.0)

title("四、系统架构与接口设计", 13)
body("系统采用“Android前端 + 后端评测服务 + 本地数据存储”的三层结构。前端负责交互、录音和展示；后端负责评测与纠错建议；本地数据库负责保存练习记录、单词库和统计数据。")
body("建议接口最少定义如下：")
bullet("POST /api/assessment/upload - 上传音频并触发评测")
bullet("GET /api/assessment/result/{id} - 获取评测结果")
bullet("GET /api/words - 获取单词列表")
bullet("GET /api/phonemes/{wordId} - 获取单词关联的音素与提示")
bullet("POST /api/practice/record - 保存一次练习记录")
bullet("GET /api/stats/{userId} - 获取学习统计")

title("五、数据库设计建议", 13)
db = doc.add_table(rows=1, cols=3)
db.style = "Table Grid"
db.alignment = WD_TABLE_ALIGNMENT.CENTER
db.autofit = False
db_headers = ["表名", "关键字段", "用途"]
db_widths = [Cm(3.2), Cm(8.3), Cm(3.1)]
for i, h in enumerate(db_headers):
    c = db.rows[0].cells[i]
    c.width = db_widths[i]
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(c, 120, 120, 120, 120)
    shade_cell(c, "D9EAF7")
    p = c.paragraphs[0]
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
    add_run(p, h, font_name="SimHei", size=10.5, bold=True)
for row_data in [
    ("user", "id, username, age_group, created_at", "用户"),
    ("word", "id, word_text, phonetic, difficulty, category, audio_path", "单词库"),
    ("phoneme_rule", "id, phoneme, mouth_hint, tongue_hint, note", "发音规则"),
    ("practice_record", "id, user_id, word_id, score, result_level, error_phonemes, record_time", "练习记录"),
    ("evaluation_result", "id, record_id, raw_result, feedback_text, created_at", "评测结果"),
    ("learning_stats", "id, user_id, total_practice, avg_score, common_error, updated_at", "统计汇总"),
]:
    cells = db.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].width = db_widths[i]
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cells[i], 120, 120, 120, 120)
        p = cells[i].paragraphs[0]
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.05)
        add_run(p, text, font_name="宋体", size=9.8)

title("六、开发顺序建议", 13)
for step in [
    "1. 先完成页面骨架和路由跳转。",
    "2. 再接入录音和标准发音播放。",
    "3. 接着对接评测接口并做结果页。",
    "4. 补充纠错文案和错误音素映射。",
    "5. 最后实现数字人展示和学习统计。",
]:
    body(step, indent=0.0)

title("七、风险与备选方案", 13)
for item in [
    "如果 3D 数字人开发成本过高，先用 2D 口型动画替代。",
    "如果评测接口不稳定，先固定在单词级评分，减少自由输入。",
    "如果儿童反馈文案太抽象，改为“张大嘴”“收一收嘴唇”这类简单提示。",
    "如果时间不足，优先保留训练闭环，不做复杂排行榜和社交功能。",
]:
    bullet(item)

title("八、交付物清单", 13)
for item in [
    "Android 可运行程序",
    "后端接口或模拟服务",
    "数据库设计文档",
    "系统架构图与流程图",
    "测试记录与演示截图",
    "论文与答辩PPT",
]:
    bullet(item)

doc.core_properties.author = "OpenAI"
doc.core_properties.title = "面向儿童英语发音的数字人口型与舌位可视化纠错系统 - 项目落地蓝图"
doc.core_properties.subject = "项目蓝图"
doc.core_properties.category = "Blueprint"

doc.save("项目落地蓝图.docx")
